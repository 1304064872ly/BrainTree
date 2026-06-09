"""
文件解析器模块
=============

本模块负责从上传的文件中提取纯文本内容。
支持的文件格式：PDF、DOCX、TXT、Markdown

功能特点：
1. PDF 解析：支持三种方式（pymupdf 优先、PyPDF2 回退、OCR 兜底）
2. DOCX 解析：同时提取段落和表格内容
3. TXT 解析：支持多种编码（UTF-8、GBK、GB2312 等）
4. 自动检测文件格式并选择合适的解析方法

与 file_converter.py 的区别：
- file_parser: 提取原始文本（用于数据库存储）
- file_converter: 转换为结构化 Markdown（用于 AI 分析）

主要类：
    FileParser: 文件解析器，提供统一的解析接口

使用示例：
    parser = FileParser()
    text = await parser.parse("document.pdf", file_bytes)
"""

# ============================================================
# 第一部分：导入依赖
# ============================================================
import io           # 内存文件流操作
import os           # 文件路径处理
import re           # 正则表达式
from typing import Optional  # 类型注解
import PyPDF2       # PDF 解析库（备选方案）
from docx import Document    # DOCX 解析库（python-docx）

# ============================================================
# 第二部分：可选依赖导入
# ============================================================
# 这些库可能未安装，使用 try-except 处理
# 如果未安装，对应的功能将不可用

# pymupdf：高性能 PDF 解析库（推荐）
# 比 PyPDF2 有更好的文本提取能力，特别是对于复杂格式的 PDF
try:
    import fitz  # pymupdf 的导入名是 fitz
    HAS_PYMUPDF = True  # 标记 pymupdf 可用
except ImportError:
    HAS_PYMUPDF = False  # pymupdf 不可用

# OCR 相关库：用于解析扫描版 PDF
# pytesseract: Tesseract OCR 的 Python 封装
# pdf2image: 将 PDF 页面转换为图片
try:
    import pytesseract  # OCR 引擎
    from pdf2image import convert_from_bytes  # PDF 转图片
    HAS_OCR = True  # 标记 OCR 可用
except ImportError:
    HAS_OCR = False  # OCR 不可用

class FileParser:
    """
    文件解析器类

    支持从 PDF、DOCX、TXT、Markdown 文件中提取纯文本内容。

    解析策略：
    - PDF: 优先使用 pymupdf，回退到 PyPDF2，最后尝试 OCR
    - DOCX: 使用 python-docx 提取段落和表格
    - TXT: 尝试多种编码（UTF-8、GBK、GB2312 等）

    属性：
        无状态属性，所有方法都是独立的

    方法：
        parse(filename, content): 解析文件的主入口
    """

    async def parse(self, filename: str, content: bytes) -> str:
        """
        解析文件内容的主入口

        根据文件扩展名自动选择合适的解析方法。
        支持的格式：.pdf, .docx, .txt, .md, .markdown

        Args:
            filename: 文件名（包含扩展名）
            content: 文件内容的字节流

        Returns:
            str: 提取的纯文本内容

        Raises:
            ValueError: 不支持的文件格式

        使用示例：
            parser = FileParser()
            with open("document.pdf", "rb") as f:
                text = await parser.parse("document.pdf", f.read())
        """
        # 提取文件扩展名并转换为小写
        ext = os.path.splitext(filename)[1].lower()

        # 根据扩展名选择解析方法
        if ext == ".pdf":
            return self._parse_pdf(content)           # PDF 文件
        elif ext == ".docx":
            return self._parse_docx(content)          # Word 文档
        elif ext == ".txt":
            return self._parse_txt(content)           # 纯文本文件
        elif ext in (".md", ".markdown"):
            return self._parse_txt(content)           # Markdown 文件按文本解析
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _parse_pdf(self, content: bytes) -> str:
        """
        解析 PDF 文件

        使用三级降级策略确保尽可能提取到文本：
        1. pymupdf（最佳）：高性能，支持复杂格式
        2. PyPDF2（备选）：纯 Python 实现，兼容性好
        3. OCR（兜底）：用于扫描版 PDF

        Args:
            content: PDF 文件的字节流

        Returns:
            str: 提取的文本内容

        注意：
        - 如果所有方法都失败，返回提示信息
        - 扫描版 PDF 需要安装 OCR 依赖
        """
        # 方法1: 使用 pymupdf 提取文本（推荐）
        # pymupdf 有更好的文本提取能力，特别是对于复杂格式的 PDF
        if HAS_PYMUPDF:
            try:
                return self._parse_pdf_with_pymupdf(content)
            except Exception as e:
                print(f"pymupdf 解析失败，尝试其他方法: {e}")

        # 方法2: 使用 PyPDF2 提取文本（备选）
        # PyPDF2 是纯 Python 实现，兼容性好但提取能力较弱
        text = ""  # 初始化 text 变量
        try:
            text = self._parse_pdf_with_pypdf2(content)
            # 如果提取的文本足够长，直接返回
            if text and len(text.strip()) > 50:
                return text
        except Exception as e:
            print(f"PyPDF2 解析失败: {e}")

        # 方法3: 使用 OCR 提取文本（扫描版 PDF）
        # OCR 需要安装 pytesseract 和 pdf2image
        if HAS_OCR:
            try:
                return self._parse_pdf_with_ocr(content)
            except Exception as e:
                print(f"OCR 解析失败: {e}")

        # 如果所有方法都失败
        if not text:
            return "[PDF文件内容为空或无法提取文本，建议使用OCR工具预处理]"
        return text

    def _parse_pdf_with_pymupdf(self, content: bytes) -> str:
        """
        使用 pymupdf 解析 PDF

        pymupdf 是一个高性能的 PDF 处理库，支持：
        - 文本提取
        - 图片提取
        - 页面操作

        Args:
            content: PDF 文件的字节流

        Returns:
            str: 提取的文本内容

        Raises:
            Exception: pymupdf 提取的文本过少时抛出异常
        """
        import fitz  # pymupdf 的导入名是 fitz

        # 从字节流打开 PDF 文档
        doc = fitz.open(stream=content, filetype="pdf")
        text_content = []  # 存储每页的文本

        # 遍历所有页面
        for page_num in range(len(doc)):
            page = doc[page_num]

            # 尝试提取文本
            text = page.get_text("text")
            if text and text.strip():
                text_content.append(text.strip())
            else:
                # 如果文本提取失败，尝试提取文本块
                # 文本块包含更多格式信息
                blocks = page.get_text("blocks")
                if blocks:
                    # blocks 是一个列表，每个元素是 (x0, y0, x1, y1, text, block_no, block_type)
                    # 第4个元素（索引4）是文本内容
                    page_text = "\n".join([block[4] for block in blocks if block[4].strip()])
                    if page_text.strip():
                        text_content.append(page_text.strip())

        # 关闭文档，释放资源
        doc.close()

        # 将所有页面的文本合并
        result = "\n\n".join(text_content)

        # 如果提取的文本太少，抛出异常尝试其他方法
        if not result or len(result.strip()) < 50:
            raise Exception("pymupdf 提取的文本过少")

        return result

    def _parse_pdf_with_pypdf2(self, content: bytes) -> str:
        """
        使用 PyPDF2 解析 PDF

        PyPDF2 是一个纯 Python 的 PDF 处理库，兼容性好。
        但文本提取能力不如 pymupdf。

        Args:
            content: PDF 文件的字节流

        Returns:
            str: 提取的文本内容
        """
        # 从字节流创建 PDF 读取器
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text_content = []  # 存储每页的文本

        # 遍历所有页面
        for page in pdf_reader.pages:
            try:
                # 提取页面文本
                text = page.extract_text()
                if text and text.strip():
                    text_content.append(text.strip())
            except Exception as e:
                # 某些页面可能提取失败，跳过继续
                print(f"警告: 解析PDF页面失败: {e}")
                continue

        # 将所有页面的文本合并
        return "\n\n".join(text_content)

    def _parse_pdf_with_ocr(self, content: bytes) -> str:
        """
        使用 OCR 解析扫描版 PDF

        对于扫描版 PDF（图片格式），需要使用 OCR 识别文字。
        流程：PDF → 图片 → OCR → 文本

        Args:
            content: PDF 文件的字节流

        Returns:
            str: OCR 识别的文本内容

        注意：
        - OCR 需要安装 Tesseract 和 pytesseract
        - 中文 OCR 需要安装中文语言包
        - OCR 速度较慢，但可以处理扫描版 PDF
        """
        print("正在使用 OCR 解析扫描版 PDF...")

        # 将 PDF 转换为图片
        # dpi=200 表示每英寸200像素，平衡质量和速度
        images = convert_from_bytes(content, dpi=200)

        text_content = []  # 存储每页的文本
        for i, image in enumerate(images):
            try:
                # 使用 pytesseract 进行 OCR
                # lang='chi_sim+eng' 表示识别中文简体和英文
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                if text and text.strip():
                    text_content.append(text.strip())
                print(f"OCR 完成第 {i+1}/{len(images)} 页")
            except Exception as e:
                # 某些页面可能识别失败，跳过继续
                print(f"OCR 第 {i+1} 页失败: {e}")
                continue

        # 将所有页面的文本合并
        result = "\n\n".join(text_content)
        if not result:
            return "[OCR 未能提取到文本内容]"
        return result

    def _parse_docx(self, content: bytes) -> str:
        """
        解析 DOCX 文件

        使用 python-docx 库解析 Word 文档，提取：
        1. 段落文本
        2. 表格内容

        Args:
            content: DOCX 文件的字节流

        Returns:
            str: 提取的文本内容

        Raises:
            Exception: DOCX 解析失败时抛出异常

        注意：
        - 只支持 .docx 格式，不支持旧版 .doc 格式
        - 表格内容会用 "|" 分隔各列
        """
        try:
            # 从字节流打开 Word 文档
            doc = Document(io.BytesIO(content))
            text_content = []  # 存储所有文本

            # 1. 提取段落内容
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # 2. 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    # 提取每行的单元格内容
                    row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if row_text:
                        # 用 " | " 分隔各列
                        text_content.append(" | ".join(row_text))

            # 将所有内容合并
            result = "\n\n".join(text_content)
            if not result:
                return "[DOCX文件内容为空或无法提取文本]"
            return result
        except Exception as e:
            raise Exception(f"DOCX 解析失败: {str(e)}")

    def _parse_txt(self, content: bytes) -> str:
        """
        解析 TXT 文件

        尝试多种编码解码文本文件，确保兼容性。
        支持的编码：UTF-8、GBK、GB2312、GB18030、Latin-1

        Args:
            content: TXT 文件的字节流

        Returns:
            str: 解码后的文本内容

        Raises:
            Exception: 所有编码都失败时抛出异常

        编码优先级：
        1. UTF-8（最常用）
        2. UTF-8 BOM（带 BOM 的 UTF-8）
        3. GBK（中文 Windows 常用）
        4. GB2312（简体中文）
        5. GB18030（扩展中文）
        6. Latin-1（西欧字符，不会失败）
        """
        try:
            # 尝试多种编码
            encodings = [
                "utf-8",       # 最常用的 Unicode 编码
                "utf-8-sig",   # 带 BOM 的 UTF-8
                "gbk",         # 中文 Windows 常用
                "gb2312",      # 简体中文
                "gb18030",     # 扩展中文
                "latin-1"      # 西欧字符（不会失败）
            ]

            for encoding in encodings:
                try:
                    result = content.decode(encoding)
                    if result and result.strip():
                        return result
                except (UnicodeDecodeError, LookupError):
                    # 编码不匹配，尝试下一个
                    continue

            # 如果所有编码都失败，使用 utf-8 并忽略错误
            result = content.decode("utf-8", errors="ignore")
            if not result or not result.strip():
                return "[TXT文件内容为空]"
            return result
        except Exception as e:
            raise Exception(f"TXT 解析失败: {str(e)}")
